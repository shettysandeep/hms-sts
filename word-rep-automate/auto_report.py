# -*- coding: utf-8 -*-

'''
Title: Automate STS report generation for ACSD Registry
Author: Sandeep Shetty
Date: Dec 10, 2021
'''

from docx import Document
import pandas as pd
import numpy as np
import re
from os.path import exists

## Excel Dataset
tables_data = pd.read_excel('ACSD Site Report Tables 1_2_3 EXCEL WORKBOOK.xlsx',sheet_name=None)

# Data for Tables 1 - 3 in the report
table1=tables_data['ACSD Table 1']
table2=tables_data['ACSD Table 2']
table3=tables_data['ACSD Table 3']

# OR LOG table for Table 4 in the report
table_data_4 = pd.read_csv('or_tables_dec14 v1_excel.csv')

### Obtain Site ID and Site Names
siteId_names = pd.read_csv('site_id_names.csv')

### Iterate over Site IDs
for index in siteId_names.index[60:]:
    SITEID, SITENAME = siteId_names.loc[index].values
    print ("in process {}{}".format(SITEID, SITENAME))
    ## Read Document Template1
    document = Document('ACSDTemplate2021 Dec15_12PM.docx')

    ## Table1 ~~~~~~~~~~~~~~~~~~~~~~
    ### Get data for Table 1
    table1['NewColumn'] = table1['Sequence Number'].astype(str).apply(lambda x: x.replace('.0','')) + '-' + table1['Form/Data Element'] + str(':')
    cols_needed = [ 'Form/Data Element', 'Sequence Number', 'NewColumn']
    all_cols=table1.columns[table1.columns.str.contains(str(SITEID))].tolist() + cols_needed
    table1A=table1[all_cols].copy()
    site_agree = 'ID_{} Agreement Rate'.format(str(SITEID))
    site_mismatch = 'ID_{} Mismatch #'.format(str(SITEID))
    ovl_ag = table1A[table1A['Form/Data Element']=='Overall Agreement Rate'][site_agree].tolist()[0]
    ovl_ag_rate = str(round(ovl_ag*100,1))+'%'
    
    ### Add data to Table 1
    #table1.head(2)
    table = document.tables[0]
    len_table_0 = len(document.tables[0].rows)
    mytext = table.rows[0].cells[1].text
    table.rows[0].cells[1].text = str(SITEID)+ ' ' + str(SITENAME) + ' '+ mytext
    for i in range(len(document.tables[0].rows)-1):
        if i > 0:

            row_label = table.rows[i].cells[0].text
            #print(row_label)
            strip_val = row_label.strip()
            if (table1A['Form/Data Element']==strip_val).any():
                #print("here")
                #site_agree = 'ID_{} Agreement Rate'.format(str(SITEID))
                temp_val = table1A[table1A['Form/Data Element']==strip_val][site_agree].tolist()[0]
                #print(temp_val)
                if str(temp_val) != "":
                    pct_ag = round(temp_val*100,1)
                    table.rows[i].cells[1].text = str(pct_ag)+'%'
                    if pct_ag>=98:
                        over_grade = "Exceeds Expectations"
                    elif (pct_ag>=90 and pct_ag<98):
                        over_grade = "Meets Expectations"
                    else:
                        over_grade = "Does Not Meet Expectations"
                else:
                    table.rows[i].cells[1].text = ''
            else:
                table.rows[i].cells[1].text = ''

    table.rows[15].cells[1].text = over_grade
    # overall_grade = table.rows[len(document.tables[0].rows)-1].cells[2].text # last row - last column 
    
    # Table 2 ~~~~~~~~~~~~~~~~~~~~~~~~
    ### Get data to Table 2
    table2['NewColumn'] = table2['Sequence Number'].astype(str).apply(lambda x: x.replace('.0','')) + '-' + table2['Form/Data Element'] + str(':')
    cols_needed_2 = [ 'Form/Data Element', 'Sequence Number', 'NewColumn']
    all_cols2=table2.columns[table2.columns.str.contains(str(SITEID))].tolist() + cols_needed_2

    crate = table2[table2['Form/Data Element']=='Complications Agreement Rate'][site_agree].tolist()[0]
    crate_rnd = round(crate*100,1)
    cratepct=str(crate_rnd)+'%'
    mrate = table2[table2['Form/Data Element']=='Mortality Agreement Rate'][site_agree].tolist()[0]
    mrate_rnd = round(mrate*100,1)
    mratepct=str(mrate_rnd)+'%'
    

    #table2.head(2)
    table = document.tables[1]
    text2=table.rows[0].cells[1].text
    table.rows[0].cells[1].text=str(SITEID)+' '+text2
    for i in range(len(document.tables[1].rows)):
        
        if i > 0:
            row_label = table.rows[i].cells[0].text
            strip_val = row_label.strip()
            if (table2['Form/Data Element']==strip_val).any():
                temp_val = table2[table2['Form/Data Element']==strip_val][site_agree].tolist()[0]
                #print(temp_val)
                if not np.isnan(temp_val):
                    table.rows[i].cells[1].text = str(round(temp_val*100,1))+'%'
                else:
                    table.rows[i].cells[1].text = ''
            else:
                table.rows[i].cells[1].text = ""

    # Adding grades to Table 2

    if crate_rnd>=98:
        table.rows[5].cells[1].text = "Meets expectations"
    elif (crate_rnd>=90 and crate_rnd<98):
        table.rows[5].cells[1].text = "Does not meet expectations"
    else:
        table.rows[5].cells[1].text = "Requires re-audit"

    if mrate_rnd>=98:
        table.rows[8].cells[1].text = "Meets expectations"
    elif (mrate_rnd>=90 and mrate_rnd<98):
        table.rows[8].cells[1].text = "Does not meet expectations"
    else:
        table.rows[8].cells[1].text = "Requires re-Audit"
        
    # Obtain Letter Grade
    comp_grade = table.rows[5].cells[1].text # Complications Grade
    mort_grade = table.rows[8].cells[1].text # Mortality Grade
    
    ### Add Data to Table 3
    table3['NewColumn'] = table3['Sequence Number'].astype(str).apply(lambda x: x.replace('.0','')) + '-' + table3['Form/Data Element'] + str(':')
    cols_needed_3 = [ 'Form/Data Element', 'Sequence Number', 'NewColumn']
    all_cols3=table3.columns[table3.columns.str.contains(str(SITEID))].tolist() + cols_needed_3
    table3A=table3[all_cols3].copy()

    table3A['Sequence Number']=table3A['Sequence Number'].astype(str).apply(lambda x: x.replace('.0',''))

    for i in range(len(document.tables[2].rows)):
        if i > 0:
            table = document.tables[2]
            row_label = table.rows[i].cells[0].text
            if re.match('\d', row_label):
                strip_val = row_label.split('-')[0]
                if (table3A['Sequence Number']==strip_val).any():
                    temp_val = table3A[table3A['Sequence Number']==strip_val][site_agree].tolist()[0]
                    temp_2 = table3A[table3A['Sequence Number']==strip_val][site_mismatch].tolist()[0]
                    table.rows[i].cells[3].text = str(round(temp_val*100,1))+'%'
                    table.rows[i].cells[2].text = str(temp_2)
                else:
                    table.rows[i].cells[3].text = ""
                    table.rows[i].cells[2].text = ""
            else:
                row_label=row_label.title()
                if (table3A['Form/Data Element']==row_label).any():
                    temp_val1 = table3A[table3A['Form/Data Element']==row_label][site_agree].tolist()[0]                
                    temp_val2 = table3A[table3A['Form/Data Element']==row_label][site_mismatch].tolist()[0]
                    table.rows[i].cells[3].text = str(round(temp_val1*100,1))+'%'
                    table.rows[i].cells[2].text = str(temp_val2)
                else:
                    table.rows[i].cells[3].text = ""
                    table.rows[i].cells[2].text = ""

    # overall_mismatches = table.rows[len(document.tables[2].rows)-1].cells[1].text
    
    # # Table 4. ~~~~~~~~~~~~~~~~~~~~~~
    ### Add Data to Table 4
    table4=document.tables[3]
    new_table=table_data_4[table_data_4['Site ID']==SITEID].copy()
    new_table = new_table[['Total Records in IQVIA', 'Matched Records in IQVIA', '%agreement']] #, 'Result']]
    #print(new_table)
    if new_table.shape[0]>0:
        for k in range(3):
            table4.rows[1].cells[k].text = str(new_table.iloc[0,k])

   # Table 5 ~~~~~~~~~~~~~~~~~~~~~~~
    ### Mismatch Table 5  
    # write a code to look for a file in the directory. 
    site_filname='ACSDSite_{}.csv'.format(str(SITEID))
    filepath='./Table 5 Mismatches for Each Site/{}'.format(site_filname)

    table5 = document.tables[4]

    if exists(filepath):
        site_table = pd.read_csv(filepath)
        site_table=site_table[['Record ID', 'Data Element', 'SequenceNumber']]
        site_tab_rows = site_table.shape[0]
        
        # adding rows
        if site_tab_rows > 0:
            n_r = site_tab_rows-len(table5.rows)
            if n_r>0:
                for i in range(n_r+1):
                    table5.add_row()
            
            for i in range(site_tab_rows): #34
                # if i > 0:
                for j in range(3):
                        # print(table5.rows[i].cells[j].text)
                    table5.rows[i+1].cells[j].text = str(site_table.iloc[i,j])

    ### Replace Site ID and Site Name
    ### Sub the Site ID and Site Name

    body_element = document._body._body
    #pb=body_element.xml

    #Site ID & SITE NAME
    body_element.xpath('./w:sdt/w:sdtContent/w:p/w:r/w:t')[13].text = str(SITEID) + ' - ' +  str(SITENAME)
    # body_element.xpath('./w:sdt/w:sdtContent/w:p/w:r/w:t')[16].text =
    # body_element.xpath('./w:p/w:r/w:t')[41].text = str(SITEID)
    for loc in [43, 54, 83]: #, 78]: #, 116, 134]:
        body_element.xpath('./w:p/w:r/w:t')[loc].text = str(SITENAME)

    body_element.xpath('./w:p/w:r/w:t')[48].text = ovl_ag_rate
    body_element.xpath('./w:p/w:r/w:t')[58].text = cratepct
    body_element.xpath('./w:p/w:r/w:t')[64].text = mratepct
    body_element.xpath('./w:p/w:r/w:t')[45].text = over_grade 
    body_element.xpath('./w:p/w:r/w:t')[62].text = " " + str(comp_grade).title()
    body_element.xpath('./w:p/w:r/w:t')[70].text = str(mort_grade).title()     
    ## Fix Conclusions
    # body_element.xpath('./w:p/w:r/w:t')[118].text =  overall_grade # Grade
    # body_element.xpath('./w:p/w:r/w:t')[121].text =  overall_mismatches # X
    # body_element.xpath('./w:p/w:r/w:t')[136].text =  # MR grademratepct


    # body_element.xpath('./w:p/w:r/w:t')[125].text = ovl_ag_rate
    # body_element.xpath('./w:p/w:r/w:t')[138].text = cratepct
    # body_element.xpath('./w:p/w:r/w:t')[141].text = mratepct
    
    
    ### Save Document
    document.save('{}-{}.docx'.format(SITEID,SITENAME))
    print("saved {}-{}".format(SITEID,SITENAME))
